"""
Build a coherent final DOCX report from the latest training artifacts.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
REPORT_PATH = BASE_DIR / "reports" / "reports.docx"


def load_json(path: str):
    return json.loads((BASE_DIR / path).read_text())


def set_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 16, "1f4e79"),
        ("Heading 2", 13, "374151"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Système intelligent de rétention client")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("1f4e79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Projet Data Science M2 - Prédiction du churn, comparaison multi-modèles, dashboard et API optionnelle")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"EFREI - Data Engineering & AI | Rapport généré le {date.today().strftime('%d/%m/%Y')}")
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def pct(value: float) -> str:
    return f"{value:.1%}"


def main() -> None:
    metadata = load_json("models/model_metadata.json")
    comparison = load_json("models/model_comparison.json")
    features = load_json("models/feature_names.json")
    imbalance = load_json("models/imbalance_study.json")
    df = pd.read_csv(BASE_DIR / "data" / "customer_churn_business_dataset.csv")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_styles(doc)
    add_title(doc)

    doc.add_heading("1. Résumé exécutif", level=1)
    doc.add_paragraph(
        "Le projet construit une solution complète de rétention client autour d'une tâche de "
        "classification binaire : prédire le churn. La solution inclut un pipeline de "
        "préparation des données, une comparaison de quatre modèles supervisés dont un MLP, "
        "une analyse d'interprétabilité, un dashboard Streamlit décisionnel et une API FastAPI optionnelle."
    )
    bullet(doc, f"Dataset utilisé : {len(df):,} clients et {df.shape[1]} colonnes, dont la cible churn.")
    bullet(doc, f"Taux de churn observé : {pct(float(df['churn'].mean()))}.")
    bullet(doc, f"Ratio de déséquilibre : {metadata['imbalance_ratio']:.2f}:1 en faveur de la classe non-churn.")
    bullet(doc, f"Modèle final sélectionné : {metadata['imbalance_strategy']} + {metadata['model_type']}.")
    bullet(doc, f"Performance test : Recall {metadata['recall']:.4f}, F1 {metadata['f1_score']:.4f}, ROC-AUC {metadata['roc_auc']:.4f}, PR-AUC {metadata['pr_auc']:.4f}.")
    bullet(doc, "Le seuil de décision est optimisé sur validation pour éviter un F1 nul et réduire les faux négatifs.")

    doc.add_heading("2. Données et problématique", level=1)
    doc.add_paragraph(
        "La variable cible est churn (0 = client conservé, 1 = client résilié). "
        "Les colonnes décrivent des dimensions démographiques, contractuelles, financières, "
        "d'engagement produit, de support client et de satisfaction."
    )
    add_table(
        doc,
        ["Indicateur", "Valeur"],
        [
            ["Nombre de clients", f"{len(df):,}"],
            ["Nombre de churners", f"{int(df['churn'].sum()):,}"],
            ["Taux de churn", pct(float(df["churn"].mean()))],
            ["Features utilisées par les modèles", str(features["n_features"])],
            ["Colonnes exclues", ", ".join(features["dropped_columns"])],
        ],
    )

    doc.add_heading("3. Préparation des données", level=1)
    bullet(doc, "Séparation train/test stratifiée : 80% entraînement, 20% test.")
    bullet(doc, "Une sous-validation stratifiée est utilisée pour ajuster les seuils de décision.")
    bullet(doc, "Le preprocessing est intégré dans des sklearn Pipeline et ColumnTransformer.")
    bullet(doc, "StandardScaler est appliqué aux variables numériques.")
    bullet(doc, "OneHotEncoder(handle_unknown='ignore') est appliqué aux variables catégorielles.")
    bullet(doc, "Les transformations sont ajustées uniquement sur le train set, ce qui limite le data leakage.")
    bullet(doc, "customer_id est exclu car il s'agit d'un identifiant technique sans valeur prédictive métier généralisable.")

    doc.add_heading("4. Gestion explicite du déséquilibre des classes", level=1)
    doc.add_paragraph(
        "Le dataset présente un déséquilibre significatif : la classe non-churn est très majoritaire. "
        "Dans ce contexte, l'accuracy seule est trompeuse. Par exemple, la baseline Logistic Regression "
        "au seuil 0.5 obtient une accuracy élevée tout en détectant presque aucun churner."
    )
    baseline = imbalance["baseline_logistic_regression"]
    add_table(
        doc,
        ["Analyse préalable", "Valeur"],
        [
            ["Classe 0 (non churn)", str(imbalance["class_distribution"]["0"])],
            ["Classe 1 (churn)", str(imbalance["class_distribution"]["1"])],
            ["Ratio majorité/minorité", f"{imbalance['imbalance_ratio']:.2f}:1"],
            ["Baseline accuracy", f"{baseline['default_accuracy']:.4f}"],
            ["Baseline recall", f"{baseline['default_recall']:.4f}"],
            ["Baseline F1", f"{baseline['default_f1']:.4f}"],
            ["Baseline PR-AUC", f"{baseline['default_pr_auc']:.4f}"],
            ["Baseline faux négatifs", str(int(baseline["default_fn"]))],
        ],
    )
    doc.add_paragraph(
        "Les métriques retenues sont Recall, F1-score, ROC-AUC et PR-AUC. Le Recall est prioritaire "
        "car un faux négatif correspond à un client à risque non détecté. La PR-AUC est ajoutée car "
        "elle est plus informative que l'accuracy lorsque la classe positive est minoritaire."
    )
    doc.add_paragraph(
        "Les méthodes testées sont : pondération des classes, Random Over-Sampling, SMOTE, "
        "Random Under-Sampling et ajustement du seuil de décision. La validation croisée utilisée "
        "est une Stratified K-Fold afin de conserver les proportions de classes dans chaque fold."
    )
    rows = []
    for row in imbalance["best_by_strategy"]:
        rows.append(
            [
                row["strategy"],
                row["model"],
                f"{row['tuned_precision']:.4f}",
                f"{row['tuned_recall']:.4f}",
                f"{row['tuned_f1']:.4f}",
                f"{row['tuned_pr_auc']:.4f}",
                f"{row['f1_threshold']:.3f}",
                str(int(row["tuned_fp"])),
                str(int(row["tuned_fn"])),
            ]
        )
    add_table(doc, ["Stratégie", "Modèle", "Precision", "Recall", "F1", "PR-AUC", "Seuil", "FP", "FN"], rows)
    doc.add_paragraph(
        f"L'approche finale retenue est {metadata['imbalance_strategy']} + {metadata['model_type']}. "
        "Elle maximise le F1 parmi les stratégies testées tout en conservant un recall élevé. "
        "Le compromis métier accepté est une hausse des faux positifs pour réduire les faux négatifs, "
        "car contacter un client faussement à risque coûte moins cher que perdre un client réellement churner."
    )

    doc.add_heading("5. Modélisation multi-algorithmes", level=1)
    rows = []
    for item in comparison["models"]:
        rows.append(
            [
                item["model"],
                f"{item['accuracy']:.4f}",
                f"{item['precision']:.4f}",
                f"{item['recall']:.4f}",
                f"{item['f1_score']:.4f}",
                f"{item['roc_auc']:.4f}",
                f"{item['threshold']:.3f}",
            ]
        )
    add_table(doc, ["Modèle", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "Seuil"], rows)
    doc.add_paragraph(
        "La comparaison montre que les modèles d'ensemble sont les plus adaptés à ce dataset tabulaire. "
        "Le modèle Deep Learning MLP est conservé dans l'analyse pour satisfaire l'exigence pédagogique "
        "et démontrer que le Deep Learning n'est pas automatiquement supérieur sur des données structurées."
    )

    doc.add_heading("6. Deep Learning et correction du F1 nul", level=1)
    doc.add_paragraph(
        "Le problème F1=0 provenait d'un seuil de décision inadapté au déséquilibre des classes : "
        "le modèle pouvait produire des probabilités mais ne classer aucun client en churn au seuil 0.5. "
        "La correction consiste à optimiser le seuil sur validation et à évaluer ensuite sur test."
    )
    mlp = next(item for item in comparison["models"] if item["model"] == "MLP")
    bullet(doc, f"MLP après correction : F1 {mlp['f1_score']:.4f}, recall {mlp['recall']:.4f}, ROC-AUC {mlp['roc_auc']:.4f}.")
    bullet(doc, "Le F1 n'est plus nul, ce qui rend la comparaison exploitable.")
    bullet(doc, "Le MLP reste inférieur aux meilleurs modèles d'ensemble, ce qui est cohérent pour un dataset tabulaire de taille modérée.")

    doc.add_heading("7. Interprétabilité", level=1)
    doc.add_paragraph(
        "L'explicabilité repose sur la permutation importance du modèle final. Cette méthode mesure "
        "la baisse de performance lorsque chaque variable est mélangée, ce qui donne une importance "
        "agnostique au modèle et plus robuste qu'une simple importance native."
    )
    rows = [
        [row["feature"], f"{row['importance']:.5f}", f"{row['std']:.5f}"]
        for row in metadata["feature_importance"][:10]
    ]
    add_table(doc, ["Feature", "Importance", "Ecart-type"], rows)
    doc.add_paragraph(
        "SHAP est documenté comme extension avancée. Dans le rendu actuel, la permutation importance "
        "constitue l'explication globale principale, directement alignée avec le modèle final."
    )

    doc.add_heading("8. Dashboard décisionnel", level=1)
    bullet(doc, "Le dashboard Streamlit affiche les KPI de churn, le revenu mensuel à risque et les segments à risque.")
    bullet(doc, "La page Prediction utilise directement le pipeline entraîné, sans formule simulée.")
    bullet(doc, "La page Modeles compare les quatre modèles et affiche le seuil de décision retenu.")
    bullet(doc, "Le dashboard ne dépend pas de l'API, conformément au choix projet ; l'API reste optionnelle.")

    doc.add_heading("9. API optionnelle", level=1)
    doc.add_paragraph(
        "Une API FastAPI est disponible comme brique d'industrialisation optionnelle. "
        "Elle expose notamment /health, /predict, /batch_predict, /model/info et des routes de comparaison. "
        "Elle n'est pas nécessaire au fonctionnement du dashboard."
    )

    doc.add_heading("10. Limites et recommandations", level=1)
    bullet(doc, "Surveiller le recall de la classe churn, car les faux négatifs ont un coût métier élevé.")
    bullet(doc, "Tester des stratégies métier de seuil selon le budget de campagne de rétention.")
    bullet(doc, "Ajouter SHAP complet en amélioration pour expliquer chaque prédiction individuelle.")
    bullet(doc, "Mettre en place un suivi de drift et un réentraînement périodique.")

    doc.add_heading("11. Conclusion", level=1)
    doc.add_paragraph(
        "Le projet répond à l'objectif principal : transformer un dataset client en système prédictif "
        "multi-modèles exploitable par un utilisateur métier. Les corrections apportées rendent les "
        "résultats cohérents avec les données actuelles, corrigent le F1 nul du MLP et remplacent la "
        "simulation dashboard par une vraie inférence du modèle final."
    )

    REPORT_PATH.parent.mkdir(exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
